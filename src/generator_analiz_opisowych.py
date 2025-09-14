#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Jul 19 12:25:54 2025

@author: adrian
"""
from tkinter import Tk, filedialog, messagebox
import pandas as pd
from datetime import datetime
import os
from qgis.core import QgsProject

project_path = QgsProject.instance().fileName()
project_directory = os.path.dirname(project_path)

# Wybór biblioteki - można przełączać między python-docx a docxtpl
USE_DOCXTPL = True  # Zmień na False, aby użyć python-docx

if USE_DOCXTPL:
    try:
        from docxtpl import DocxTemplate
        print("Używanie docxtpl (nowsza biblioteka z szablonami Jinja2)")
    except ImportError:
        print("Instaluję docxtpl...")
        os.system("pip install docxtpl")
        from docxtpl import DocxTemplate
else:
    from docx import Document
    print("Używanie python-docx (klasyczna biblioteka)")


def replace_placeholder_docx(doc, replacements):
    """Funkcja dla python-docx"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, val in replacements.items():
                placeholder = f"{{{key}}}"  # Szuka {klucz} w dokumencie
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(val))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_docx(cell, replacements)

def generate_document():
    # Ukrycie głównego okna Tkinter
    root = Tk()
    root.withdraw()
    
    # Ścieżka początkowa
    initial_dir = "/home/adrian/Documents/JXPROJEKT/analizy"
    template_dir = "/home/adrian/Documents/JXPROJEKT/analizaWZ_szablony"
    
    try:
        # 1. Wybór szablonu Word
        print("📄 Wybierz szablon dokumentu...")
        template_path = filedialog.askopenfilename(
            title="Wybierz szablon dokumentu Word",
            filetypes=[("Dokumenty Word", "*.docx")],
            initialdir=template_dir
        )
        
        if not template_path:
            print("❌ Nie wybrano szablonu. Przerywanie.")
            return
        
        print(f"✅ Wybrano szablon: {os.path.basename(template_path)}")
        
        # 2. Wybór plików Excel z danymi
        print("📊 Wybierz pliki Excel z danymi...")
        file_paths = filedialog.askopenfilenames(
            title="Wybierz pliki: dane_dzialki_inwestora oraz results_excel_export",
            filetypes=[("Pliki Excel", "*.xlsx")],
            initialdir=initial_dir,
            multiple=True
        )
        
        if len(file_paths) != 2:
            messagebox.showerror("Błąd", "Musisz wybrać dokładnie dwa pliki Excel.")
            return
        
        print(f"✅ Wybrano pliki: {[os.path.basename(f) for f in file_paths]}")
        
        # 3. Wczytanie danych z plików Excel
        print("🔄 Wczytywanie danych...")
        excel_data = pd.read_excel(file_paths[0], sheet_name="do_eksportu")
        excel_data_results = pd.read_excel(file_paths[1])
        
        # Dodanie dzisiejszej daty
        dzisiaj = datetime.now().date().strftime("%d.%m.%Y")
        data_biezaca = {'nazwa_pola': 'today', 'wartosc': dzisiaj}
        
        # 4. Przetwarzanie danych dzialki inwestora
        for i in range(len(excel_data)):
            try: 
                excel_data.loc[i, 'wartosc'] = round(excel_data.loc[i, 'wartosc'], 2)
            except TypeError:
                pass
      
        
        excel_data = pd.concat([excel_data, pd.DataFrame([data_biezaca])], ignore_index=True)
        excel_data.loc[1, 'wartosc'] = excel_data.loc[1, 'wartosc'].date().strftime("%d.%m.%Y")
        dane_dzialki = excel_data.set_index('nazwa_pola')['wartosc'].to_dict()
        
        # 5. Przetwarzanie wyników
        for i in range(len(excel_data_results)):
            if isinstance(excel_data_results.loc[i, 'wartosc'], (int, float)):
                excel_data_results.loc[i, 'wartosc'] = round(excel_data_results.loc[i, 'wartosc'], 2)
        
        results_data = excel_data_results.set_index('nazwa_pola')['wartosc'].to_dict()
        
        # 6. Łączenie danych
        replacements = {**dane_dzialki, **results_data}
        znak_sprawy = excel_data.loc[0, 'wartosc']
        
        print(f"🔧 Znaleziono {len(replacements)} zmiennych do zastąpienia")
        
        # 7. Generowanie dokumentu
        print("📝 Generowanie dokumentu...")
        
        if USE_DOCXTPL:
            # Użycie docxtpl (nowsza metoda z Jinja2)
            doc = DocxTemplate(template_path)
            doc.render(replacements)
        else:
            # Użycie python-docx (klasyczna metoda)
            doc = Document(template_path)
            replace_placeholder_docx(doc, replacements)
        
        # 8. Zapisanie dokumentu
        katalog_sprawy = file_paths[1].split('/')[-2]
        output_path = f"/home/adrian/Documents/JXPROJEKT/analizy/{katalog_sprawy}/{znak_sprawy}-analiza.docx"
        
        # Utworzenie katalogu jeśli nie istnieje
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc.save(output_path)
        
        print(f'✅ Gotowe! Nowy plik zapisany jako "{znak_sprawy}-analiza.docx"')
        print(f'📁 Lokalizacja: {output_path}')
        
        # Opcjonalne wyświetlenie listy zastąpionych zmiennych
        print("\n📋 Zastąpione zmienne:")
        for key, value in list(replacements.items())[:10]:  # Pokaż pierwsze 10
            print(f"   {key}: {value}")
        if len(replacements) > 10:
            print(f"   ... i {len(replacements) - 10} więcej")
            
    except Exception as e:
        print(f"❌ Wystąpił błąd: {e}")
        messagebox.showerror("Błąd", f"Wystąpił błąd: {e}")
    
    finally:
        root.destroy()


generate_document()